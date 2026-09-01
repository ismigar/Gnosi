"""Document and media adapters for LLM Wiki source extraction."""

from __future__ import annotations

import html
import logging
import re
import shutil
import subprocess
import tempfile
from collections.abc import Callable
from pathlib import Path
from typing import Any


Segment = dict[str, Any]
TextExtractor = Callable[[Path], str]
SegmentsExtractor = Callable[[Path], list[Segment]]
TemporaryRoot = Callable[[], Path]


def extract_pdf(
    path: Path,
    *,
    run_tesseract: TextExtractor,
    temporary_root: TemporaryRoot,
    logger: logging.Logger,
) -> list[Segment]:
    """Extract page and paragraph locators from a PDF, with bounded OCR fallback."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    segments: list[Segment] = []
    pdfium: Any = None
    for page_number, page in enumerate(reader.pages, start=1):
        text = str(page.extract_text() or "").strip()
        if len(re.sub(r"\s+", "", text)) < 30 and shutil.which("tesseract"):
            try:
                if pdfium is None:
                    import pypdfium2  # type: ignore[import-untyped]  # Third-party adapter lacks py.typed.

                    pdfium = pypdfium2.PdfDocument(str(path))
                with tempfile.NamedTemporaryFile(suffix=".png", dir=temporary_root()) as tmp:
                    image = pdfium[page_number - 1].render(scale=2.0).to_pil()
                    image.save(tmp.name)
                    text = run_tesseract(Path(tmp.name))
            except Exception as error:
                logger.warning("llm_wiki PDF OCR failed on page %s: %s", page_number, error)
        for paragraph_number, paragraph in enumerate(split_paragraphs(text), start=1):
            segments.append(
                {
                    "text": paragraph,
                    "locator": {"page": page_number, "paragraph": paragraph_number},
                }
            )
    return segments


def extract_docx(path: Path) -> list[Segment]:
    """Extract paragraphs with their most recent document heading."""
    from docx import Document

    document = Document(str(path))
    segments: list[Segment] = []
    heading = ""
    paragraph_number = 0
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            heading = text
            continue
        paragraph_number += 1
        segments.append(
            {
                "text": text,
                "locator": {"section": heading, "paragraph": paragraph_number},
            }
        )
    return segments


def extract_epub(path: Path) -> list[Segment]:
    """Extract ordered paragraph locators from EPUB document items."""
    from bs4 import BeautifulSoup
    from ebooklib import ITEM_DOCUMENT, epub  # type: ignore[import-untyped]  # Third-party adapter lacks py.typed.

    book = epub.read_epub(str(path))
    segments: list[Segment] = []
    chapter_number = 0
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        chapter_number += 1
        soup = BeautifulSoup(item.get_content(), "html.parser")
        title_node = soup.find(["h1", "h2", "title"])
        chapter = title_node.get_text(" ", strip=True) if title_node else item.get_name()
        for paragraph_number, node in enumerate(
            soup.find_all(["p", "li", "blockquote"]),
            start=1,
        ):
            text = node.get_text(" ", strip=True)
            if text:
                segments.append(
                    {
                        "text": text,
                        "locator": {
                            "chapter": chapter,
                            "chapter_number": chapter_number,
                            "paragraph": paragraph_number,
                        },
                    }
                )
    return segments


def extract_html(raw_html: str) -> list[Segment]:
    """Extract readable HTML blocks with heading context."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(raw_html, "html.parser")
    for node in soup(["script", "style", "noscript", "svg"]):
        node.decompose()
    segments: list[Segment] = []
    heading = ""
    paragraph_number = 0
    for node in soup.find_all(["h1", "h2", "h3", "p", "li", "blockquote"]):
        text = html.unescape(node.get_text(" ", strip=True))
        if not text:
            continue
        if node.name in {"h1", "h2", "h3"}:
            heading = text
            continue
        paragraph_number += 1
        segments.append(
            {
                "text": text,
                "locator": {"section": heading, "paragraph": paragraph_number},
            }
        )
    return segments


def extract_text_file(path: Path) -> list[Segment]:
    """Extract paragraph and line locators from a UTF-8-compatible text file."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    return paragraph_segments(raw, locator_prefix="lines")


def paragraph_segments(raw: str, *, locator_prefix: str) -> list[Segment]:
    """Split text without losing stable line-range locators."""
    segments: list[Segment] = []
    line_cursor = 1
    for paragraph in re.split(r"\n\s*\n+", str(raw or "")):
        text = paragraph.strip()
        if not text:
            line_cursor += paragraph.count("\n") + 1
            continue
        line_count = text.count("\n") + 1
        segments.append(
            {
                "text": " ".join(line.strip() for line in text.splitlines() if line.strip()),
                "locator": {
                    "kind": locator_prefix,
                    "line_start": line_cursor,
                    "line_end": line_cursor + line_count - 1,
                },
            }
        )
        line_cursor += line_count + 1
    return segments


def extract_image(
    path: Path,
    *,
    run_tesseract: TextExtractor,
) -> list[Segment]:
    """Extract OCR paragraphs with visual locators."""
    text = run_tesseract(path)
    return [
        {"text": paragraph, "locator": {"image": path.name, "paragraph": index}}
        for index, paragraph in enumerate(split_paragraphs(text), start=1)
    ]


def run_tesseract(path: Path, *, extraction_error: type[RuntimeError]) -> str:
    """Run OCR using the languages available on the current host."""
    binary = shutil.which("tesseract")
    if not binary:
        raise extraction_error("Tesseract is not installed")
    languages = available_tesseract_languages(binary)
    requested = [language for language in ("cat", "spa", "eng", "fra") if language in languages]
    command = [binary, str(path), "stdout"]
    if requested:
        command.extend(["-l", "+".join(requested)])
    process = subprocess.run(
        command,
        capture_output=True,
        text=True,
        timeout=180,
        check=False,
    )
    if process.returncode != 0:
        raise extraction_error(process.stderr.strip() or "Tesseract failed")
    return process.stdout.strip()


def available_tesseract_languages(binary: str) -> set[str]:
    """Return installed OCR languages, or an empty set when probing fails."""
    try:
        process = subprocess.run(
            [binary, "--list-langs"],
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
        return {line.strip() for line in process.stdout.splitlines()[1:] if line.strip()}
    except Exception:
        return set()


def extract_audio(path: Path) -> list[Segment]:
    """Extract timed transcript segments from an audio file."""
    from backend.services.transcription import transcribe

    result = transcribe(str(path))
    raw_segments = result.get("segments") or []
    if not isinstance(raw_segments, list):
        return []
    return [
        {
            "text": str(item.get("text") or "").strip(),
            "locator": {
                "start": float(item.get("start") or 0),
                "end": float(item.get("end") or 0),
            },
        }
        for item in raw_segments
        if isinstance(item, dict) and str(item.get("text") or "").strip()
    ]


def extract_video(
    path: Path,
    *,
    extract_audio_segments: SegmentsExtractor,
    run_tesseract: TextExtractor,
    temporary_root: TemporaryRoot,
    logger: logging.Logger,
) -> list[Segment]:
    """Combine timed transcription with bounded visual keyframe OCR."""
    segments = extract_audio_segments(path)
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg or not shutil.which("tesseract"):
        return segments
    with tempfile.TemporaryDirectory(
        prefix="gnosi-llm-wiki-frames-",
        dir=temporary_root(),
    ) as temporary_directory:
        pattern = str(Path(temporary_directory) / "frame-%04d.jpg")
        command = [
            ffmpeg,
            "-hide_banner",
            "-loglevel",
            "error",
            "-i",
            str(path),
            "-vf",
            "fps=1/60,scale=1280:-2",
            "-frames:v",
            "40",
            pattern,
        ]
        process = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=600,
            check=False,
        )
        if process.returncode != 0:
            logger.warning("llm_wiki keyframe extraction failed: %s", process.stderr.strip())
            return segments
        frames = sorted(Path(temporary_directory).glob("frame-*.jpg"))
        for frame_index, frame in enumerate(frames, start=1):
            try:
                text = run_tesseract(frame)
            except Exception:
                continue
            for paragraph in split_paragraphs(text):
                segments.append(
                    {
                        "text": paragraph,
                        "locator": {
                            "start": float((frame_index - 1) * 60),
                            "frame": frame_index,
                            "visual": True,
                        },
                    }
                )
    return sorted(segments, key=_segment_start)


def _segment_start(segment: Segment) -> float:
    locator = segment.get("locator")
    if not isinstance(locator, dict):
        return 0.0
    return float(locator.get("start") or 0)


def split_paragraphs(text: str) -> list[str]:
    """Normalize paragraphs while retaining their source order."""
    raw = str(text or "").strip()
    if not raw:
        return []
    return [
        " ".join(piece.split())
        for piece in re.split(r"\n\s*\n+|(?<=\.)\s*\n+", raw)
        if " ".join(piece.split())
    ]


__all__ = [
    "Segment",
    "available_tesseract_languages",
    "extract_audio",
    "extract_docx",
    "extract_epub",
    "extract_html",
    "extract_image",
    "extract_pdf",
    "extract_text_file",
    "extract_video",
    "paragraph_segments",
    "run_tesseract",
    "split_paragraphs",
]

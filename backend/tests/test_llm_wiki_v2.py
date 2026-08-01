"""Regression tests for the version 2 LLM Wiki contracts."""
from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from backend.services import (
    llm_wiki,
    llm_wiki_config,
    llm_wiki_extractors,
    llm_wiki_indices,
    llm_wiki_storage,
    llm_wiki_suggestions,
)


def _origin(text: str, *, label: str = "Source", order: int = 0) -> dict:
    return llm_wiki_extractors._finalize_origin({  # noqa: SLF001
        "kind": "text",
        "label": label,
        "source_url": "",
        "input_order": order,
        "segments": [{"text": text, "locator": {"line_start": 1, "line_end": 1}}],
    })


def test_v1_configuration_migrates_brain_and_reference_source():
    config = llm_wiki_config.normalize_config(
        {"target_table": "brain-1", "configured": True},
        reference_table_id="references-1",
    )

    assert config["version"] == 2
    assert config["brain_table_id"] == "brain-1"
    assert config["target_table"] == "brain-1"
    assert config["source_tables"] == [{
        "table_id": "references-1",
        "title_property_id": "",
        "attachment_property_ids": [],
        "url_property_ids": [],
        "language_property_id": "",
        "include_body": True,
        "relation_property_id": "",
        "dimension_mappings": {},
    }]
    assert config["ui_locale"] == "en"


def test_default_brain_schema_and_names_are_english():
    from backend.api.vault_routes import _brain_schema

    assert llm_wiki_config.normalize_config({})["ui_locale"] == "en"
    assert _brain_schema()[0] == ("note_type", "Note type", "select")
    assert all(role != "legacy_source" for role, _name, _type in _brain_schema())
    assert all(name != "Sources" for _role, name, _type in _brain_schema())


def test_visible_note_type_uses_existing_semantic_option():
    prop = {
        "name": "Tipus de nota",
        "type": "select",
        "config": {
            "options": [
                {"name": "Nota índex"},
                {"name": "Nota permanent"},
                {"name": "Nota de lectura"},
            ],
        },
    }
    config = {"ui_locale": "ca"}

    assert llm_wiki_config.note_type_value("reading", config, prop) == "Nota de lectura"
    assert llm_wiki_config.note_type_value("index", config, prop) == "Nota índex"
    assert llm_wiki_config.note_type_value("permanent", config, prop) == "Nota permanent"
    assert llm_wiki_config.note_type_value("system", config, prop) == "Nota de sistema"
    assert llm_wiki_config.metadata_note_type(
        {"Tipus de nota": "Nota permanent"},
    ) == "permanent"


def test_legacy_brain_property_ids_are_recovered_from_configured_mappings(monkeypatch):
    from backend.api import vault_routes

    source = {
        "id": "resources",
        "properties": [
            {"id": "source-area", "name": "Àrees", "type": "relation"},
            {"id": "source-project", "name": "Projecte", "type": "relation"},
        ],
    }
    brain = {
        "id": "brain",
        "properties": [
            {"name": "Tipus de nota", "type": "select"},
            {"name": "Àrea", "type": "relation"},
            {"name": "Projecte", "type": "relation"},
        ],
    }
    config = {
        "brain_roles": {"note_type": "note-type-id", "areas": "area-id"},
        "index_field_ids": ["area-id", "project-id"],
        "source_tables": [{
            "table_id": "resources",
            "dimension_mappings": {
                "area-id": {
                    "mode": "source",
                    "source_property_id": "source-area",
                },
                "project-id": {
                    "mode": "source",
                    "source_property_id": "source-project",
                },
            },
        }],
    }
    monkeypatch.setattr(
        vault_routes,
        "_table_by_id",
        lambda table_id: source if table_id == "resources" else None,
    )

    hints = vault_routes._brain_property_id_hints(config, brain)

    assert hints["tipusdenota"] == "note-type-id"
    assert hints["area"] == "area-id"
    assert hints["projecte"] == "project-id"


def test_resource_brain_view_uses_contextual_singular_source_filter():
    from backend.api.vault_routes import _normalize_brain_source_view

    view = {
        "filters": [
            {
                "field": "Font",
                "operator": "equals",
                "value": "http://localhost:5173/vault/page/resource-uid",
            },
            {"field": "Tipus de nota", "value": "Nota de lectura"},
        ],
    }

    assert _normalize_brain_source_view(view, "Font", {"Font", "Fonts"}) is True
    assert view["filters"] == [
        {"field": "Font", "value": "this"},
        {"field": "Tipus de nota", "value": "Nota de lectura"},
    ]
    assert _normalize_brain_source_view(view, "Font", {"Font", "Fonts"}) is False


def test_brain_source_relation_merges_and_removes_plural_field(monkeypatch, tmp_path: Path):
    from backend.api import vault_routes

    registry = {
        "tables": [
            {
                "id": "brain",
                "name": "Cervell",
                "properties": [
                    {
                        "name": "Font",
                        "type": "relation",
                        "relation_database_id": "resources",
                        "cardinality": "many-to-one",
                    },
                    {
                        "name": "Fonts",
                        "type": "relation",
                        "relation_database_id": "resources",
                        "cardinality": "many-to-one",
                    },
                ],
            },
            {"id": "resources", "name": "Recursos", "properties": []},
        ],
        "views": [],
    }
    note_path = tmp_path / "note.md"
    note_path.write_text("placeholder", encoding="utf-8")
    saved_page: dict = {}
    monkeypatch.setattr(vault_routes, "load_registry", lambda: registry)
    monkeypatch.setattr(vault_routes, "save_registry", lambda _registry: None)
    monkeypatch.setattr(
        vault_routes,
        "_get_pages_for_table",
        lambda table_id: [SimpleNamespace(path=note_path)] if table_id == "brain" else [],
    )
    monkeypatch.setattr(
        vault_routes,
        "parse_frontmatter",
        lambda _raw, _path: (
            {
                "title": "Reading note",
                "Font": ["[[Resource|resource-1]]"],
                "Fonts": [
                    "[[Resource|resource-1]]",
                    "[[Other resource|resource-2]]",
                ],
            },
            "Body",
        ),
    )
    monkeypatch.setattr(
        vault_routes,
        "save_page_md",
        lambda _path, metadata, body: saved_page.update(
            {"metadata": metadata, "body": body},
        ),
    )
    monkeypatch.setattr(vault_routes, "register_page_in_index", lambda _path: None)

    relation_id = vault_routes.ensure_brain_source_relation(
        "brain",
        "resources",
        "ca",
    )

    source_properties = [
        prop
        for prop in registry["tables"][0]["properties"]
        if prop.get("relation_database_id") == "resources"
    ]
    assert relation_id
    assert len(source_properties) == 1
    assert source_properties[0]["name"] == "Font"
    assert source_properties[0]["id"] == relation_id
    assert saved_page["metadata"]["Font"] == [
        "[[Resource|resource-1]]",
        "[[Other resource|resource-2]]",
    ]
    assert "Fonts" not in saved_page["metadata"]


def test_process_title_prefers_configured_source_title_over_uid(tmp_path: Path):
    from backend.api.vault_routes import _llm_wiki_source_title

    path = tmp_path / "acd052f1-9788-5036-bf5e-4e945806e15d.md"
    metadata = {
        "title": "acd052f1-9788-5036-bf5e-4e945806e15d",
        "Title": "Distinció entre afirmacions i judicis",
    }
    source_table = {
        "properties": [
            {"id": "title-field", "name": "Title", "type": "title"},
        ],
    }
    source_config = {"title_property_id": "title-field"}

    assert _llm_wiki_source_title(
        metadata,
        path,
        source_table,
        source_config,
    ) == "Distinció entre afirmacions i judicis"


def test_existing_reading_and_permanent_notes_follow_source_contract():
    from backend.api.vault_routes import _normalize_brain_page_contract

    brain = {
        "properties": [
            {
                "id": "note-type",
                "name": "Tipus de nota",
                "type": "select",
                "config": {
                    "options": [
                        {"name": "Nota índex"},
                        {"name": "Nota permanent"},
                        {"name": "Nota de lectura"},
                    ],
                },
            },
            {
                "id": "source",
                "name": "Font",
                "aliases": ["Fonts"],
                "type": "relation",
                "relation_database_id": "resources",
            },
        ],
    }
    config = {
        "ui_locale": "ca",
        "brain_roles": {"note_type": "note-type"},
        "source_tables": [{
            "table_id": "resources",
            "relation_property_id": "source",
        }],
    }
    reading = {
        "title": "Atomic idea",
        "note_type": "lectura",
        "llm_wiki_source_table_id": "resources",
        "llm_wiki_resource_id": "resource-1",
        "llm_wiki_resource_title": "resource-1",
        "Fonts": ["[[resource-1|resource-1]]"],
        "Tipus de nota": "lectura",
    }
    permanent = {
        "title": "Manual idea",
        "note_type": "permanent",
        "Font": ["[[Resource title|resource-1]]"],
        "Fonts": ["[[Resource title|resource-1]]"],
        "Tipus de nota": "Nota permanent",
    }

    assert _normalize_brain_page_contract(
        reading,
        config,
        brain,
        {("resources", "resource-1"): "Resource title"},
    )
    assert reading["Tipus de nota"] == "Nota de lectura"
    assert reading["llm_wiki_resource_title"] == "Resource title"
    assert reading["Font"] == ["[[Resource title|resource-1]]"]
    assert "Fonts" not in reading

    assert _normalize_brain_page_contract(permanent, config, brain, {}) is True
    assert "Font" not in permanent
    assert "Fonts" not in permanent


def test_source_detection_handles_different_schemas_and_ai_fallback():
    brain = {
        "id": "brain",
        "properties": [
            {"id": "brain-area", "name": "Àrees", "type": "multi_select"},
            {"id": "brain-tag", "name": "Etiquetes", "type": "multi_select"},
        ],
    }
    first = {
        "id": "papers",
        "properties": [
            {"id": "paper-name", "name": "Nom", "type": "title"},
            {"id": "paper-file", "name": "PDF", "type": "files"},
            {"id": "paper-url", "name": "Enllaç", "type": "url"},
            {"id": "paper-area", "name": "Àrees", "type": "select"},
        ],
    }
    second = {
        "id": "videos",
        "properties": [
            {"id": "video-title", "name": "Title", "type": "title"},
            {"id": "video-url", "name": "URL", "type": "url"},
        ],
    }

    detected_first = llm_wiki_config.auto_detect_source(
        first,
        brain,
        ["brain-area", "brain-tag"],
    )
    detected_second = llm_wiki_config.auto_detect_source(
        second,
        brain,
        ["brain-area", "brain-tag"],
    )

    assert detected_first["title_property_id"] == "paper-name"
    assert detected_first["attachment_property_ids"] == ["paper-file"]
    assert detected_first["url_property_ids"] == ["paper-url"]
    assert detected_first["dimension_mappings"]["brain-area"] == {
        "mode": "source",
        "source_property_id": "paper-area",
        "fixed_value": None,
    }
    assert detected_first["dimension_mappings"]["brain-tag"]["mode"] == "ai"
    assert detected_second["dimension_mappings"]["brain-area"]["mode"] == "ai"
    assert detected_second["dimension_mappings"]["brain-tag"]["mode"] == "ai"


def test_long_sources_are_chunked_completely_and_keep_stable_evidence_ids():
    text = "".join(str(index % 10) for index in range(31_337))
    origin = _origin(text)

    chunks = llm_wiki_extractors.chunk_origins([origin], max_chars=4_000)
    chunk_text = "".join(
        segment["text"]
        for chunk in chunks
        for segment in chunk["segments"]
    )

    assert len(chunks) == 8
    assert chunk_text == text
    assert {segment["id"] for chunk in chunks for segment in chunk["segments"]} == {
        origin["segments"][0]["id"]
    }


def test_duplicate_content_keeps_one_origin_and_all_aliases():
    first = _origin("The same normalized evidence.", label="attachment.pdf", order=0)
    second = _origin("The same normalized evidence.", label="https://example.test/a", order=1)
    second["kind"] = "url"

    deduplicated = llm_wiki_extractors._deduplicate_origins([first, second])  # noqa: SLF001

    assert len(deduplicated) == 1
    assert deduplicated[0]["label"] == "attachment.pdf"
    assert deduplicated[0]["aliases"] == [{
        "kind": "url",
        "label": "https://example.test/a",
        "source_url": "",
        "input_order": 1,
    }]


def test_document_extractors_keep_structured_locators(tmp_path: Path):
    from docx import Document
    from ebooklib import epub
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    text_path = tmp_path / "fixture.md"
    text_path.write_text("# Heading\n\nFirst paragraph.\n\nSecond paragraph.", encoding="utf-8")

    docx_path = tmp_path / "fixture.docx"
    document = Document()
    document.add_heading("Section one", level=1)
    document.add_paragraph("DOCX paragraph evidence.")
    document.save(docx_path)

    epub_path = tmp_path / "fixture.epub"
    book = epub.EpubBook()
    book.set_identifier("fixture")
    book.set_title("Fixture")
    book.set_language("en")
    chapter = epub.EpubHtml(title="Chapter one", file_name="chapter.xhtml", lang="en")
    chapter.content = "<h1>Chapter one</h1><p>EPUB paragraph evidence.</p>"
    book.add_item(chapter)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(epub_path, book)

    pdf_path = tmp_path / "fixture.pdf"
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
    })
    resources = DictionaryObject({
        NameObject("/Font"): DictionaryObject({
            NameObject("/F1"): writer._add_object(font),  # noqa: SLF001
        }),
    })
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 72 720 Td (PDF paragraph evidence.) Tj ET")
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(content)  # noqa: SLF001
    with pdf_path.open("wb") as handle:
        writer.write(handle)

    text_segments = llm_wiki_extractors._extract_text_file(text_path)  # noqa: SLF001
    docx_segments = llm_wiki_extractors._extract_docx(docx_path)  # noqa: SLF001
    epub_segments = llm_wiki_extractors._extract_epub(epub_path)  # noqa: SLF001
    pdf_segments = llm_wiki_extractors._extract_pdf(pdf_path)  # noqa: SLF001

    assert [segment["locator"]["line_start"] for segment in text_segments] == [1, 3, 5]
    assert docx_segments == [{
        "text": "DOCX paragraph evidence.",
        "locator": {"section": "Section one", "paragraph": 1},
    }]
    assert epub_segments[0]["locator"]["chapter"] == "Chapter one"
    assert epub_segments[0]["text"] == "EPUB paragraph evidence."
    assert pdf_segments[0]["locator"] == {"page": 1, "paragraph": 1}
    assert "PDF paragraph evidence" in pdf_segments[0]["text"]


def test_multimedia_extractors_preserve_time_and_visual_locators(monkeypatch, tmp_path: Path):
    from backend.services import transcription

    audio = tmp_path / "fixture.wav"
    audio.write_bytes(b"fixture")
    monkeypatch.setattr(
        transcription,
        "transcribe",
        lambda _path: {
            "segments": [
                {"start": 1.25, "end": 3.5, "text": "Timed transcript evidence."},
            ],
        },
    )
    assert llm_wiki_extractors._extract_audio(audio) == [{  # noqa: SLF001
        "text": "Timed transcript evidence.",
        "locator": {"start": 1.25, "end": 3.5},
    }]

    image = tmp_path / "fixture.png"
    image.write_bytes(b"fixture")
    monkeypatch.setattr(
        llm_wiki_extractors,
        "_run_tesseract",
        lambda _path: "First visual paragraph.\n\nSecond visual paragraph.",
    )
    assert llm_wiki_extractors._extract_image(image) == [  # noqa: SLF001
        {
            "text": "First visual paragraph.",
            "locator": {"image": "fixture.png", "paragraph": 1},
        },
        {
            "text": "Second visual paragraph.",
            "locator": {"image": "fixture.png", "paragraph": 2},
        },
    ]


def test_public_url_helpers_detect_podcast_media_and_block_local_ssrf(monkeypatch):
    from backend.agent import web_context

    html = b"""
        <html><head><meta property="og:audio" content="/episode.mp3"></head>
        <body><p>Episode notes.</p></body></html>
    """
    assert (
        llm_wiki_extractors._embedded_media_url(  # noqa: SLF001
            html,
            "https://podcast.example/show/1",
        )
        == "https://podcast.example/episode.mp3"
    )
    monkeypatch.setattr(
        web_context,
        "is_public_http_url",
        lambda _url: (False, "private network"),
    )
    with pytest.raises(llm_wiki_extractors.ExtractionError, match="Unsafe URL blocked"):
        llm_wiki_extractors._download_public_url("http://127.0.0.1/private")  # noqa: SLF001


def test_plan_validation_requires_exact_evidence_and_preserves_source_order():
    first = _origin("First grounded idea appears here.", label="A", order=0)
    second = _origin("Second grounded idea appears later.", label="B", order=1)
    first_segment = first["segments"][0]["id"]
    second_segment = second["segments"][0]["id"]
    plans = [
        (
            {"segments": second["segments"]},
            {"notes": [{
                "title": "Second idea",
                "body_md": "One idea.",
                "source_segment_id": second_segment,
                "dimensions": {"area": ["Research"]},
                "citations": [{
                    "segment_id": second_segment,
                    "quote": "Second grounded idea",
                }],
            }]},
        ),
        (
            {"segments": first["segments"]},
            {"notes": [
                {
                    "title": "First idea",
                    "body_md": "One idea.",
                    "source_segment_id": first_segment,
                    "citations": [{
                        "segment_id": first_segment,
                        "quote": "First grounded idea",
                    }],
                },
                {
                    "title": "Invented idea",
                    "body_md": "Unsupported.",
                    "source_segment_id": first_segment,
                    "citations": [{
                        "segment_id": first_segment,
                        "quote": "This text is not in the source",
                    }],
                },
            ]},
        ),
    ]
    dimensions = [{
        "field_id": "area",
        "multiple": False,
        "by_label": {"research": "Research"},
    }]

    notes, warnings = llm_wiki._validate_and_reduce_plans(  # noqa: SLF001
        plans,
        [first, second],
        dimensions,
    )

    assert [note["title"] for note in notes] == ["First idea", "Second idea"]
    assert [note["position"] for note in notes] == [1, 2]
    assert notes[1]["dimensions"] == {"area": "Research"}
    assert len({note["managed_key"] for note in notes}) == 2
    assert warnings == ["Ungrounded model note skipped: Invented idea"]


def test_extended_citation_contains_snapshot_segment_and_locator():
    citation = llm_wiki._render_citations(  # noqa: SLF001
        [{
            "quote": "Exact evidence",
            "segment_id": "seg-1",
            "snapshot_id": "snapshot-1",
            "origin_id": "origin-1",
            "locator": {"page": 7, "paragraph": 3, "start": 12.5, "end": 18.0},
        }],
        "Resource",
        "resource-1",
        "source-table-1",
    )

    assert "gnosi-cite:?" in citation
    assert "res=resource-1" in citation
    assert "table=source-table-1" in citation
    assert "snapshot=snapshot-1" in citation
    assert "segment=seg-1" in citation
    assert "page=7" in citation
    assert "paragraph=3" in citation
    assert "start=12.5" in citation
    assert "[[Resource|" not in citation


def test_relation_dimension_values_accept_normalized_source_ids():
    prop = {"type": "relation"}
    options = [{
        "id": "area-1",
        "label": "Area title",
        "value": "[[Area title|area-1]]",
    }]

    assert llm_wiki._canonical_dimension_value(  # noqa: SLF001
        prop,
        ["area-1"],
        options,
    ) == ["[[Area title|area-1]]"]


def test_source_dimensions_override_generated_values():
    merged = llm_wiki._effective_dimensions(  # noqa: SLF001
        {"tags": ["AI tag"], "area": ["AI area"]},
        {"tags": ["Source tag"], "project": ["Source project"]},
    )

    assert merged == {
        "tags": ["Source tag"],
        "area": ["AI area"],
        "project": ["Source project"],
    }


def test_redundant_source_link_is_removed_from_managed_citation():
    body = (
        "> Evidence — [p. 1](gnosi-cite:?res=resource-1&page=1) · "
        "[[Resource|resource-1]]\n"
    )

    assert llm_wiki_indices._remove_redundant_source_links(body) == (  # noqa: SLF001
        "> Evidence — [p. 1](gnosi-cite:?res=resource-1&page=1)\n"
    )


def test_index_rebuild_synchronizes_existing_source_dimensions(monkeypatch, tmp_path):
    source_page = SimpleNamespace(
        id="resource-1",
        title="Resource",
        metadata={"Area source": ["area-1"]},
    )
    reading_path = tmp_path / "reading.md"
    reading_path.write_text("placeholder", encoding="utf-8")
    reading_page = SimpleNamespace(
        id="reading-1",
        title="Reading",
        path=str(reading_path),
        metadata={
            "note_type": "lectura",
            "llm_wiki_managed": True,
            "llm_wiki_source_table_id": "sources",
            "llm_wiki_resource_id": "resource-1",
        },
    )
    tables = {
        "brain": {
            "properties": [
                {"id": "area", "name": "Area", "type": "relation"},
            ],
        },
        "sources": {
            "properties": [
                {"id": "source-area", "name": "Area source", "type": "relation"},
            ],
        },
    }
    saved = []
    monkeypatch.setattr(llm_wiki_indices, "_table", lambda table_id: tables[table_id])
    monkeypatch.setattr(
        llm_wiki_indices,
        "_brain_pages",
        lambda table_id: [reading_page] if table_id == "brain" else [source_page],
    )
    monkeypatch.setattr(
        llm_wiki,
        "_dimension_context",
        lambda *_args: ({"area": ["[[Area|area-1]]"]}, []),
    )
    monkeypatch.setattr(
        llm_wiki_indices,
        "_read_page",
        lambda _path: (
            {
                **reading_page.metadata,
                "Area": ["[[Old area|old-area]]"],
            },
            "> Evidence — [p. 1](gnosi-cite:?res=resource-1&page=1) · "
            "[[Resource|resource-1]]\n",
        ),
    )
    monkeypatch.setattr(
        llm_wiki_indices,
        "_save_existing_page",
        lambda path, metadata, body: saved.append((path, metadata, body)),
    )
    config = {
        "source_tables": [{
            "table_id": "sources",
            "dimension_mappings": {
                "area": {
                    "mode": "source",
                    "source_property_id": "source-area",
                },
            },
        }],
    }

    count = llm_wiki_indices.sync_source_dimensions("brain", config)

    assert count == 1
    assert saved[0][1]["Area"] == ["[[Area|area-1]]"]
    assert "[[Resource|resource-1]]" not in saved[0][2]


def test_managed_blocks_preserve_manual_content():
    original = "# Manual introduction\n\nManual conclusion.\n"
    first = llm_wiki_indices._replace_managed_block(  # noqa: SLF001
        original,
        "general",
        "- [[First index]]",
    )
    second = llm_wiki_indices._replace_managed_block(  # noqa: SLF001
        first,
        "general",
        "- [[Updated index]]",
    )

    assert "# Manual introduction" in second
    assert "Manual conclusion." in second
    assert "[[First index]]" not in second
    assert second.count("<!-- gnosi:llm-wiki:start general -->") == 1
    assert "[[Updated index]]" in second
    assert (
        "<!-- gnosi:llm-wiki:start general -->\n\n"
        "- [[Updated index]]\n\n"
        "<!-- gnosi:llm-wiki:end general -->"
    ) in second


def test_resource_index_is_flat_and_preserves_appearance_order(monkeypatch):
    captured = []
    readings = [
        SimpleNamespace(
            id="note-2",
            title="Second",
            metadata={
                "llm_wiki_origin_order": 1,
                "llm_wiki_origin_label": "URL",
                "llm_wiki_resource_title": "Resource",
                "Posició": 2,
                "Àrees": ["Research"],
            },
        ),
        SimpleNamespace(
            id="note-1",
            title="First",
            metadata={
                "llm_wiki_origin_order": 0,
                "llm_wiki_origin_label": "Attachment",
                "llm_wiki_resource_title": "Resource",
                "Posició": 1,
                "Àrees": ["Research"],
            },
        ),
    ]
    monkeypatch.setattr(
        llm_wiki_indices,
        "_upsert_managed_page",
        lambda *args, **kwargs: captured.append((args, kwargs)) or {
            "id": "resource-index",
            "title": args[1],
        },
    )

    result = llm_wiki_indices._upsert_resource_index(  # noqa: SLF001
        "brain",
        "sources",
        "resource-1",
        readings,
        {"relation_property_id": "relation-source"},
        {"ui_locale": "en", "index_field_ids": ["area"], "brain_roles": {}},
        {
            "area": {"id": "area", "name": "Àrees", "type": "multi_select"},
            "relation-source": {
                "id": "relation-source",
                "name": "Source · Papers",
                "type": "relation",
            },
        },
    )

    body = captured[0][0][4]
    metadata = captured[0][0][5]
    assert result["title"] == "Index · Resource"
    assert body == "1. [[note-1|First]]\n2. [[note-2|Second]]"
    assert "Attachment" not in body
    assert "URL" not in body
    assert "Resource:" not in body
    assert metadata["Source · Papers"] == ["[[Resource|resource-1]]"]
    assert metadata["Àrees"] == ["Research"]


def test_index_sorting_accepts_legacy_position_ranges():
    assert llm_wiki_indices._sortable_integer("254-255") == 254  # noqa: SLF001
    assert llm_wiki_indices._sortable_integer("unknown") == 0  # noqa: SLF001


def test_dimension_index_links_readings_and_manual_permanents_separately(monkeypatch):
    captured = []
    reading = SimpleNamespace(
        id="reading-1",
        title="Reading idea",
        metadata={
            "note_type": "lectura",
            "Àrees": ["Research"],
            "llm_wiki_resource_title": "Resource",
            "llm_wiki_origin_order": 0,
            "Posició": 1,
        },
    )
    permanent = SimpleNamespace(
        id="permanent-1",
        title="Manual synthesis",
        metadata={"note_type": "permanent", "Àrees": ["Research"]},
    )
    monkeypatch.setattr(
        llm_wiki_indices,
        "_table",
        lambda _table_id: {
            "properties": [{"id": "area", "name": "Àrees", "type": "multi_select"}],
        },
    )
    monkeypatch.setattr(
        llm_wiki_indices,
        "_upsert_managed_page",
        lambda *args, **kwargs: captured.append((args, kwargs)) or {
            "id": "dimension-index",
            "title": args[1],
        },
    )

    pages = llm_wiki_indices._rebuild_dimension_indexes(  # noqa: SLF001
        "brain",
        {"id": "area", "name": "Àrees", "type": "multi_select"},
        [reading],
        [permanent],
        {"ui_locale": "en", "brain_roles": {}},
    )

    body = captured[0][0][4]
    assert len(pages) == 1
    assert "[[reading-1|Reading idea]]" in body
    assert "[[permanent-1|Manual synthesis]]" in body
    assert body.index("[[reading-1|Reading idea]]") < body.index("## Manual permanent notes")
    assert body.index("## Manual permanent notes") < body.index("[[permanent-1|Manual synthesis]]")


def test_generated_wikilinks_keep_ids_as_targets_and_titles_as_aliases():
    assert llm_wiki_indices._wikilink("note-1", "Readable title") == "[[note-1|Readable title]]"  # noqa: SLF001


def test_rebuildable_search_vectors_support_hybrid_ranking():
    query = llm_wiki_indices.search_vector("distributed knowledge graph")
    close = llm_wiki_indices.search_vector("knowledge graph")
    unrelated = llm_wiki_indices.search_vector("cooking sourdough bread")

    assert len(query) == 192
    assert llm_wiki_indices.vector_similarity(query, close) > 0.3
    assert (
        llm_wiki_indices.vector_similarity(query, close)
        > llm_wiki_indices.vector_similarity(query, unrelated)
    )


def test_permanent_note_acceptance_is_disabled():
    with pytest.raises(RuntimeError, match="cannot create permanent notes"):
        llm_wiki_suggestions.accept_suggestion("proposal-1", "brain-1")


def test_connection_queue_is_the_canonical_graph_overlay(monkeypatch, tmp_path: Path):
    from backend.api import vault_routes
    from backend.services.graph_service import GraphService

    vault = tmp_path / "vault"
    config_root = vault / ".gnosi"
    config_root.mkdir(parents=True)
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {"VAULT": vault, "GNOSI_CONFIG": config_root}[key],
    )
    (config_root / "llm_wiki_suggestions.json").write_text(
        '{"suggestions":[{"id":"proposal-1","member_ids":["a","b"],'
        '"title":"Shared concern"}]}',
        encoding="utf-8",
    )
    GraphService._graph_cache = {"stale": {"nodes": [], "edges": []}}
    GraphService._last_graph_time = {"stale": 1.0}

    assert llm_wiki_suggestions.sync_graph_mirror() == 1

    assert llm_wiki_suggestions.list_graph_edges() == [{
        "source": "a",
        "target": "b",
        "kind": "suggestion",
        "similarity": 90,
        "reason": "Shared concern",
        "suggestion_id": "proposal-1",
    }]
    assert GraphService._graph_cache == {}
    assert GraphService._last_graph_time == {}
    assert '"llm_wiki": "proposal-1"' in (vault / "suggestions.json").read_text(
        encoding="utf-8",
    )


def test_matching_checkpoint_resumes_writing_without_another_llm_call(monkeypatch, tmp_path: Path):
    origin = _origin("A grounded idea that was already planned.")
    segment_id = origin["segments"][0]["id"]
    plan = {
        "summary": "Stored summary",
        "notes": [{
            "title": "Stored atomic idea",
            "body_md": "One idea.",
            "managed_key": "stable-key",
            "position": 1,
            "origin_id": origin["origin_id"],
            "origin_order": 0,
            "origin_label": origin["label"],
            "source_segment_id": segment_id,
            "citations": [{
                "segment_id": segment_id,
                "quote": "A grounded idea",
                "locator": origin["segments"][0]["locator"],
                "origin_id": origin["origin_id"],
                "origin_label": origin["label"],
                "snapshot_id": "snapshot-1",
                "source_url": "",
            }],
            "dimensions": {},
        }],
    }
    config = {
        "version": 2,
        "brain_table_id": "brain",
        "source_tables": [{"table_id": "sources"}],
        "index_field_ids": [],
        "brain_roles": {},
    }
    applied = []
    manifests = []
    monkeypatch.setattr(llm_wiki.llm_wiki_config, "load_config", lambda: config)
    monkeypatch.setattr(
        llm_wiki.llm_wiki_extractors,
        "extract_resource_sources",
        lambda *_args, **_kwargs: ([origin], []),
    )
    monkeypatch.setattr(
        llm_wiki.llm_wiki_storage,
        "save_snapshot",
        lambda *_args, **_kwargs: {
            "snapshot_id": "snapshot-1",
            "kind": "text",
            "label": "Source",
            "source_url": "",
        },
    )
    monkeypatch.setattr(llm_wiki, "_load_brain_index", lambda *_args: [])
    monkeypatch.setattr(llm_wiki, "_dimension_context", lambda *_args: ({}, []))
    monkeypatch.setattr(
        llm_wiki,
        "_apply_plan",
        lambda next_plan, *_args, **_kwargs: (
            applied.append(next_plan)
            or {"created": [], "created_ids": [], "updated": ["Stored atomic idea"]}
        ),
    )
    monkeypatch.setattr(llm_wiki.llm_wiki_storage, "load_manifest", lambda *_args: {})
    monkeypatch.setattr(
        llm_wiki.llm_wiki_storage,
        "save_manifest",
        lambda *_args: manifests.append(_args[-1]),
    )
    monkeypatch.setattr(
        "backend.agent.factory.generate_text",
        lambda *_args, **_kwargs: pytest.fail("A matching write checkpoint must not call the LLM"),
    )

    report = llm_wiki.process_resource(
        "resource-1",
        "Resource",
        {"table_id": "sources"},
        "",
        "brain",
        tmp_path,
        source_table_id="sources",
        source_table={"id": "sources", "properties": []},
        source_config={"table_id": "sources"},
        resume_checkpoint={
            "plan": plan,
            "origin_hashes": [origin["content_hash"]],
            "model": "stored-model",
        },
    )

    assert applied == [plan]
    assert report["model"] == "stored-model"
    assert report["updated"] == ["Stored atomic idea"]
    assert manifests[0]["managed_keys"] == ["stable-key"]


def test_jobs_snapshots_and_manifests_are_persistent(monkeypatch, tmp_path: Path):
    from backend.api import vault_routes

    vault = tmp_path / "vault"
    local_data = tmp_path / "local-data"
    config_root = vault / ".gnosi"
    vault.mkdir()
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {
            "VAULT": vault,
            "LOCAL_DATA": local_data,
            "GNOSI_CONFIG": config_root,
        }[key],
    )
    monkeypatch.setattr(llm_wiki_storage, "_JOBS", {})
    monkeypatch.setattr(llm_wiki_storage, "_RUNNING_BY_RESOURCE", {})

    job = llm_wiki_storage.create_job("sources", "resource-1")
    llm_wiki_storage.update_job(job["job_id"], phase="writing", progress=70)
    llm_wiki_storage.finish_job(job["job_id"], phase="done", created=["note-1"])
    status = llm_wiki_storage.get_job_status("resource-1", "sources")

    origin = _origin("Immutable source paragraph.")
    snapshot = llm_wiki_storage.save_snapshot("sources", "resource-1", origin)
    evidence = llm_wiki_storage.load_evidence(
        "resource-1",
        snapshot["snapshot_id"],
        origin["segments"][0]["id"],
    )
    llm_wiki_storage.save_manifest(
        "sources",
        "resource-1",
        {"job_id": job["job_id"], "managed_note_ids": ["note-1"]},
    )

    assert status["phase"] == "done"
    assert status["created"] == ["note-1"]
    assert evidence is not None
    assert evidence["segment"]["text"] == "Immutable source paragraph."
    assert llm_wiki_storage.load_manifest("sources", "resource-1")["managed_note_ids"] == ["note-1"]


def test_managed_page_state_lives_in_synced_sidecar(monkeypatch, tmp_path: Path):
    from backend.api import vault_routes

    vault = tmp_path / "vault"
    config_root = vault / ".gnosi"
    vault.mkdir()
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {
            "VAULT": vault,
            "GNOSI_CONFIG": config_root,
        }[key],
    )
    llm_wiki_storage._PAGE_STATE_CACHE.clear()  # noqa: SLF001
    metadata = {
        "id": "note-1",
        "title": "Portable note",
        "table_id": "brain",
        "Note type": "Reading note",
        "note_type": "lectura",
        "llm_wiki_managed": True,
        "llm_wiki_key": "stable-key",
        "llm_wiki_resource_id": "resource-1",
    }

    portable = llm_wiki_storage.prepare_managed_markdown(metadata)
    restored = llm_wiki_storage.merge_page_metadata(portable, "note-1")

    assert portable == {
        "id": "note-1",
        "title": "Portable note",
        "table_id": "brain",
        "Note type": "Reading note",
    }
    assert restored["note_type"] == "lectura"
    assert restored["llm_wiki_key"] == "stable-key"
    assert restored["llm_wiki_resource_id"] == "resource-1"
    assert (
        llm_wiki_storage.page_state_path("note-1")
        == config_root / "llm_wiki" / "pages" / "note-1.json"
    )


def test_legacy_managed_frontmatter_migrates_without_body_changes(
    monkeypatch,
    tmp_path: Path,
):
    from backend.api import vault_routes

    vault = tmp_path / "vault"
    config_root = vault / ".gnosi"
    note_path = vault / "Reading.md"
    vault.mkdir()
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {
            "VAULT": vault,
            "GNOSI_CONFIG": config_root,
        }[key],
    )
    monkeypatch.setattr(vault_routes, "register_page_in_index", lambda _path: None)
    vault_routes.save_page_md(
        note_path,
        {
            "id": "reading-1",
            "title": "Reading",
            "table_id": "brain",
            "Note type": "Reading note",
            "note_type": "lectura",
            "llm_wiki_managed": True,
            "llm_wiki_key": "stable-key",
        },
        "Manual body stays unchanged.\n",
    )
    monkeypatch.setattr(
        llm_wiki_indices,
        "_brain_pages",
        lambda _table_id: [
            SimpleNamespace(
                id="reading-1",
                title="Reading",
                path=str(note_path),
                metadata={},
            ),
        ],
    )
    llm_wiki_storage._PAGE_STATE_CACHE.clear()  # noqa: SLF001

    migrated = llm_wiki_indices.migrate_managed_frontmatter("brain")
    portable, body = vault_routes.parse_frontmatter(
        note_path.read_text(encoding="utf-8"),
        note_path,
    )

    assert migrated == 1
    assert "note_type" not in portable
    assert not any(key.startswith("llm_wiki_") for key in portable)
    assert portable["Note type"] == "Reading note"
    assert body.strip() == "Manual body stays unchanged."
    assert llm_wiki_storage.load_page_state("reading-1")["llm_wiki_key"] == "stable-key"


def test_new_reading_note_writes_portable_frontmatter_and_managed_sidecar(
    monkeypatch,
    tmp_path: Path,
):
    from backend.api import vault_routes

    vault = tmp_path / "vault"
    brain_dir = vault / "Brain"
    config_root = vault / ".gnosi"
    brain_dir.mkdir(parents=True)
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {
            "VAULT": vault,
            "GNOSI_CONFIG": config_root,
        }[key],
    )
    monkeypatch.setattr(vault_routes, "_get_pages_for_table", lambda _table_id: [])
    monkeypatch.setattr(
        vault_routes,
        "_resolve_table_folder_from_metadata",
        lambda _metadata: brain_dir,
    )
    monkeypatch.setattr(
        vault_routes,
        "_get_unique_filepath",
        lambda folder, title, suffix: folder / f"{title}{suffix}",
    )
    monkeypatch.setattr(
        vault_routes,
        "_table_by_id",
        lambda _table_id: {
            "properties": [
                {
                    "id": "note-type",
                    "name": "Note type",
                    "type": "select",
                    "options": [{"name": "Reading note"}],
                },
            ],
        },
    )
    saved = []
    monkeypatch.setattr(
        vault_routes,
        "save_page_md",
        lambda path, metadata, body: saved.append((path, metadata, body)),
    )
    monkeypatch.setattr(vault_routes, "register_page_in_index", lambda _path: None)
    llm_wiki_storage._PAGE_STATE_CACHE.clear()  # noqa: SLF001
    config = {
        "ui_locale": "en",
        "brain_roles": {"note_type": "note-type"},
        "source_tables": [],
        "index_field_ids": [],
    }
    plan = {
        "notes": [{
            "title": "Atomic idea",
            "type": "concepte",
            "body_md": "Portable body.",
            "tags": [],
            "managed_key": "stable-key",
            "position": 1,
            "origin_id": "origin-1",
            "origin_order": 0,
            "origin_label": "Source",
            "source_segment_id": "segment-1",
            "dimensions": {},
            "citations": [],
        }],
    }

    result = llm_wiki._apply_plan(  # noqa: SLF001
        plan,
        "resource-1",
        "Resource",
        "brain",
        source_table_id="sources",
        config=config,
    )

    assert result["created"] == ["Atomic idea"]
    assert len(saved) == 1
    written_metadata = saved[0][1]
    assert written_metadata["Note type"] == "Reading note"
    assert "note_type" not in written_metadata
    assert not any(key.startswith("llm_wiki_") for key in written_metadata)
    page_id = written_metadata["id"]
    state = llm_wiki_storage.load_page_state(page_id)
    assert state["note_type"] == "lectura"
    assert state["llm_wiki_key"] == "stable-key"
    assert state["llm_wiki_resource_id"] == "resource-1"


def test_v2_config_http_contract_uses_a_disposable_vault(monkeypatch, tmp_path: Path):
    from backend.api import vault_routes
    from backend.services.workspace_service import WorkspaceContext

    vault = tmp_path / "vault"
    config_root = vault / ".gnosi"
    local_data = tmp_path / "local-data"
    config_root.mkdir(parents=True)
    brain = {
        "id": "brain",
        "name": "Brain",
        "properties": [
            {"id": "note-type", "name": "Note type", "type": "select"},
            {
                "id": "area",
                "name": "Areas",
                "type": "multi_select",
                "options": [{"name": "Research"}],
            },
        ],
    }
    source = {
        "id": "papers",
        "name": "Papers",
        "properties": [
            {"id": "title", "name": "Title", "type": "title"},
            {"id": "file", "name": "File", "type": "files"},
            {"id": "url", "name": "URL", "type": "url"},
            {"id": "area-source", "name": "Areas", "type": "select"},
        ],
    }
    tables = {"brain": brain, "papers": source}
    monkeypatch.setattr(
        vault_routes,
        "get_p",
        lambda key: {
            "VAULT": vault,
            "GNOSI_CONFIG": config_root,
            "LOCAL_DATA": local_data,
        }[key],
    )
    monkeypatch.setattr(vault_routes, "_table_by_id", lambda table_id: tables.get(table_id))
    monkeypatch.setattr(vault_routes, "_get_pages_for_table", lambda _table_id: [])
    monkeypatch.setattr(vault_routes, "_ensure_default_db_group", lambda: None)
    monkeypatch.setattr(vault_routes, "ensure_brain_table_schema", lambda *_args: 0)
    monkeypatch.setattr(
        vault_routes,
        "_infer_brain_roles",
        lambda _table: {"note_type": "note-type", "areas": "area"},
    )
    monkeypatch.setattr(
        vault_routes,
        "ensure_brain_source_relation",
        lambda _brain, table_id, *_args: f"relation-{table_id}",
    )
    monkeypatch.setattr(vault_routes, "_load_plugins_state", lambda: {"disabled": []})
    monkeypatch.setattr(vault_routes, "_llm_wiki_enabled", lambda _state: True)
    monkeypatch.setattr(llm_wiki_indices, "ensure_system_pages", lambda *_args: {})

    app = FastAPI()
    app.include_router(vault_routes.router, prefix="/api/vault")
    context = WorkspaceContext("test", "test-user", "owner", vault, ["read", "write"])
    app.dependency_overrides[vault_routes.get_workspace_context] = lambda: context
    client = TestClient(app)
    payload = {
        "version": 2,
        "ui_locale": "en",
        "brain_table_id": "brain",
        "source_tables": [{
            "table_id": "papers",
            "title_property_id": "title",
            "attachment_property_ids": ["file"],
            "url_property_ids": ["url"],
            "language_property_id": "",
            "include_body": False,
            "relation_property_id": "",
            "dimension_mappings": {
                "area": {
                    "mode": "source",
                    "source_property_id": "area-source",
                    "fixed_value": None,
                },
            },
        }],
        "index_field_ids": ["area"],
    }

    response = client.put("/api/vault/llm-wiki/config", json=payload)
    assert response.status_code == 200, response.text
    saved = response.json()
    assert saved["config"]["version"] == 2
    assert saved["config"]["ui_locale"] == "en"
    assert saved["config"]["source_tables"][0]["relation_property_id"] == "relation-papers"
    assert saved["validation"]["valid"] is True
    assert saved["index_options"]["area"] == [{"label": "Research", "value": "Research"}]
    assert saved["capabilities"]["modules"]["yt_dlp"] is True

    loaded = client.get("/api/vault/llm-wiki/config")
    assert loaded.status_code == 200
    assert loaded.json()["config"] == saved["config"]

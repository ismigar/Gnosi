"""Ordered, provenance-preserving source extraction for LLM Wiki.

Every adapter returns the same JSON-friendly shape: an origin contains ordered
segments, and every segment has a stable id, text, and source-specific locator.
No adapter truncates source content.  LLM-sized chunking happens afterwards.
"""
from __future__ import annotations

import asyncio
import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Iterable, Optional
from urllib.parse import unquote, urljoin, urlparse

import requests

from backend.config.logger_config import get_logger
from backend.services import llm_wiki_config

logger = get_logger(__name__)

_HTTP_TIMEOUT = (10, 90)
_USER_AGENT = "Gnosi-LLM-Wiki/2.0"
_MAX_DOWNLOAD_BYTES = int(os.environ.get("GNOSI_LLM_WIKI_MAX_DOWNLOAD_MB", "500")) * 1024 * 1024

PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst"}
HTML_EXTENSIONS = {".html", ".htm"}
DOCX_EXTENSIONS = {".docx"}
EPUB_EXTENSIONS = {".epub"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".heic"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus", ".webm"}
VIDEO_EXTENSIONS = {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".mpeg", ".mpg", ".webm"}


class ExtractionError(RuntimeError):
    """A source exists but cannot be converted into readable segments."""


def capability_report() -> dict[str, Any]:
    """Report optional runtime capabilities for Settings diagnostics."""
    modules = {}
    for module in ("pypdfium2", "docx", "ebooklib", "yt_dlp", "faster_whisper"):
        try:
            __import__(module)
            modules[module] = True
        except Exception:
            modules[module] = False
    binaries = {
        "ffmpeg": bool(shutil.which("ffmpeg")),
        "ffprobe": bool(shutil.which("ffprobe")),
        "tesseract": bool(shutil.which("tesseract")),
    }
    tesseract_languages = (
        sorted(_available_tesseract_languages(shutil.which("tesseract") or ""))
        if binaries["tesseract"]
        else []
    )
    required_ocr_languages = {"cat", "spa", "eng", "fra"}
    return {
        "modules": modules,
        "binaries": binaries,
        "supported_extensions": sorted(
            PDF_EXTENSIONS
            | TEXT_EXTENSIONS
            | HTML_EXTENSIONS
            | DOCX_EXTENSIONS
            | EPUB_EXTENSIONS
            | IMAGE_EXTENSIONS
            | AUDIO_EXTENSIONS
            | VIDEO_EXTENSIONS
        ),
        "streaming": modules["yt_dlp"] and binaries["ffmpeg"],
        "ocr": binaries["tesseract"],
        "ocr_languages": tesseract_languages,
        "ocr_missing_languages": sorted(required_ocr_languages - set(tesseract_languages)),
        "transcription": modules["faster_whisper"],
    }


def extract_resource_sources(
    metadata: dict,
    body: str,
    vault_root: Path,
    source_table: dict,
    source_config: dict,
) -> tuple[list[dict[str, Any]], list[str]]:
    """Extract every configured attachment followed by every configured URL.

    Exact duplicate content is represented once and records all equivalent
    origins under ``aliases``.  The page body is a fallback when no configured
    attachment or URL yielded readable content.
    """
    props_by_id = {
        str(prop.get("id") or ""): prop
        for prop in source_table.get("properties") or []
        if isinstance(prop, dict) and prop.get("id")
    }
    raw_inputs: list[tuple[str, str]] = []
    for prop_id in source_config.get("attachment_property_ids") or []:
        prop = props_by_id.get(str(prop_id))
        for value in _values_for_property(metadata, prop):
            raw_inputs.append(("attachment", value))
    for prop_id in source_config.get("url_property_ids") or []:
        prop = props_by_id.get(str(prop_id))
        for value in _values_for_property(metadata, prop):
            if value.lower().startswith(("http://", "https://")):
                raw_inputs.append(("url", value))

    origins: list[dict[str, Any]] = []
    warnings: list[str] = []
    for input_order, (input_kind, value) in enumerate(raw_inputs):
        try:
            if input_kind == "url":
                extracted = _extract_url(value, input_order)
            else:
                path = _resolve_attachment_path(value, Path(vault_root))
                if not path:
                    raise ExtractionError(f"Attachment not found or outside configured roots: {value}")
                _materialize(path)
                extracted = _extract_path(path, input_order, label=Path(path).name)
            origins.extend(extracted)
        except Exception as exc:  # noqa: BLE001
            message = f"{input_kind} {value}: {exc}"
            logger.warning("llm_wiki source extraction skipped: %s", message)
            warnings.append(message)

    if not origins and source_config.get("include_body", False) and str(body or "").strip():
        segments = _paragraph_segments(str(body), locator_prefix="body")
        if segments:
            origins.append(_finalize_origin({
                "kind": "body",
                "label": str(metadata.get("title") or "Resource body"),
                "source_url": "",
                "input_order": len(raw_inputs),
                "segments": segments,
            }))

    return _deduplicate_origins(origins), warnings


def chunk_origins(
    origins: list[dict[str, Any]],
    *,
    max_chars: int = 12000,
) -> list[dict[str, Any]]:
    """Create complete ordered LLM chunks without dropping any segment."""
    chunks: list[dict[str, Any]] = []
    for origin in origins:
        current: list[dict[str, Any]] = []
        current_chars = 0
        for segment in origin.get("segments") or []:
            text = str(segment.get("text") or "")
            if current and current_chars + len(text) > max_chars:
                chunks.append(_chunk(origin, current, len(chunks)))
                current = []
                current_chars = 0
            # A single very long paragraph is split with line-range continuity.
            if len(text) > max_chars:
                if current:
                    chunks.append(_chunk(origin, current, len(chunks)))
                    current = []
                    current_chars = 0
                for part_index, start in enumerate(range(0, len(text), max_chars)):
                    part = dict(segment)
                    # Keep the evidence id stable: the persisted snapshot stores
                    # the complete paragraph under this id. The part number is a
                    # processing locator only.
                    part["id"] = segment["id"]
                    part["text"] = text[start:start + max_chars]
                    part["locator"] = {
                        **(segment.get("locator") or {}),
                        "part": part_index + 1,
                    }
                    chunks.append(_chunk(origin, [part], len(chunks)))
                continue
            current.append(segment)
            current_chars += len(text)
        if current:
            chunks.append(_chunk(origin, current, len(chunks)))
    return chunks


def _chunk(origin: dict[str, Any], segments: list[dict[str, Any]], index: int) -> dict[str, Any]:
    return {
        "id": f"chunk-{index + 1}",
        "origin_id": origin["origin_id"],
        "origin_order": origin.get("input_order", 0),
        "origin_label": origin.get("label") or origin.get("kind"),
        "kind": origin.get("kind"),
        "snapshot_id": origin.get("snapshot_id"),
        "segments": segments,
    }


def _values_for_property(metadata: dict, prop: Optional[dict]) -> list[str]:
    if not prop:
        return []
    names = [str(prop.get("name") or ""), str(prop.get("id") or "")]
    raw = next(
        (
            metadata.get(name)
            for name in names
            if name and metadata.get(name) not in (None, "", [], {})
        ),
        None,
    )
    values = raw if isinstance(raw, list) else ([] if raw in (None, "") else [raw])
    out: list[str] = []
    for value in values:
        if isinstance(value, dict):
            value = value.get("url") or value.get("path") or value.get("name") or ""
        cleaned = str(value or "").strip()
        if cleaned.startswith("[[") and cleaned.endswith("]]"):
            cleaned = cleaned[2:-2].split("|", 1)[0].strip()
        if cleaned:
            out.append(cleaned)
    return out


def _resolve_attachment_path(raw: str, vault_root: Path) -> Optional[Path]:
    value = unquote(str(raw or "").strip())
    if value.lower().startswith("file://"):
        value = unquote(urlparse(value).path)
    try:
        from backend.api.vault_routes import _reroot_attachment_under_current_host

        rerooted = _reroot_attachment_under_current_host(value)
        if rerooted and rerooted.exists():
            return rerooted.resolve()
    except Exception:
        pass

    root = vault_root.resolve()
    candidate = Path(value)
    target = (candidate if candidate.is_absolute() else root / candidate).resolve()
    if target != root and root not in target.parents:
        return None
    return target if target.exists() else None


def _materialize(path: Path) -> None:
    from backend.services.files_provider import get_files_provider

    provider = get_files_provider()
    if not provider.is_online_only(path):
        return
    try:
        ok = asyncio.run(provider.materialize(path))
    except RuntimeError:
        # Extraction runs in a worker thread, but keep a safe fallback for tests
        # that invoke it from an event-loop owner.
        result: list[bool] = []

        def _runner() -> None:
            result.append(asyncio.run(provider.materialize(path)))

        import threading

        thread = threading.Thread(target=_runner, daemon=True)
        thread.start()
        thread.join()
        ok = bool(result and result[0])
    if not ok:
        raise ExtractionError(f"Cloud file could not be materialized: {path.name}")


def _extract_path(path: Path, input_order: int, *, label: str = "") -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        kind, segments = "pdf", _extract_pdf(path)
    elif suffix in DOCX_EXTENSIONS:
        kind, segments = "docx", _extract_docx(path)
    elif suffix in EPUB_EXTENSIONS:
        kind, segments = "epub", _extract_epub(path)
    elif suffix in HTML_EXTENSIONS:
        kind, segments = "html", _extract_html(path.read_text(encoding="utf-8", errors="replace"))
    elif suffix in TEXT_EXTENSIONS:
        kind, segments = suffix.lstrip("."), _extract_text_file(path)
    elif suffix in IMAGE_EXTENSIONS:
        kind, segments = "image", _extract_image(path)
    elif suffix in VIDEO_EXTENSIONS:
        kind, segments = "video", _extract_video(path)
    elif suffix in AUDIO_EXTENSIONS:
        kind, segments = "audio", _extract_audio(path)
    else:
        guessed = mimetypes.guess_type(path.name)[0] or ""
        if guessed.startswith("text/"):
            kind, segments = "text", _extract_text_file(path)
        else:
            raise ExtractionError(f"Unsupported source format: {suffix or guessed or path.name}")
    if not segments:
        raise ExtractionError(f"No readable content found in {path.name}")
    return [_finalize_origin({
        "kind": kind,
        "label": label or path.name,
        "source_url": "",
        "input_order": input_order,
        "segments": segments,
    })]


def _extract_pdf(path: Path) -> list[dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    segments: list[dict[str, Any]] = []
    pdfium = None
    for page_number, page in enumerate(reader.pages, start=1):
        text = str(page.extract_text() or "").strip()
        if len(re.sub(r"\s+", "", text)) < 30 and shutil.which("tesseract"):
            try:
                if pdfium is None:
                    import pypdfium2

                    pdfium = pypdfium2.PdfDocument(str(path))
                with tempfile.NamedTemporaryFile(suffix=".png", dir=_temporary_root()) as tmp:
                    image = pdfium[page_number - 1].render(scale=2.0).to_pil()
                    image.save(tmp.name)
                    text = _run_tesseract(Path(tmp.name))
            except Exception as exc:  # noqa: BLE001
                logger.warning("llm_wiki PDF OCR failed on page %s: %s", page_number, exc)
        paragraphs = _split_paragraphs(text)
        for paragraph_number, paragraph in enumerate(paragraphs, start=1):
            segments.append({
                "text": paragraph,
                "locator": {"page": page_number, "paragraph": paragraph_number},
            })
    return segments


def _extract_docx(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    segments = []
    heading = ""
    paragraph_number = 0
    for paragraph in doc.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            heading = text
            continue
        paragraph_number += 1
        segments.append({
            "text": text,
            "locator": {"section": heading, "paragraph": paragraph_number},
        })
    return segments


def _extract_epub(path: Path) -> list[dict[str, Any]]:
    from bs4 import BeautifulSoup
    from ebooklib import ITEM_DOCUMENT, epub

    book = epub.read_epub(str(path))
    segments = []
    chapter_number = 0
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        chapter_number += 1
        soup = BeautifulSoup(item.get_content(), "html.parser")
        title_node = soup.find(["h1", "h2", "title"])
        chapter = title_node.get_text(" ", strip=True) if title_node else item.get_name()
        for paragraph_number, node in enumerate(soup.find_all(["p", "li", "blockquote"]), start=1):
            text = node.get_text(" ", strip=True)
            if text:
                segments.append({
                    "text": text,
                    "locator": {
                        "chapter": chapter,
                        "chapter_number": chapter_number,
                        "paragraph": paragraph_number,
                    },
                })
    return segments


def _extract_html(raw_html: str) -> list[dict[str, Any]]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(raw_html, "html.parser")
    for node in soup(["script", "style", "noscript", "svg"]):
        node.decompose()
    segments = []
    heading = ""
    paragraph_number = 0
    for node in soup.find_all(["h1", "h2", "h3", "p", "li", "blockquote"]):
        text = html.unescape(node.get_text(" ", strip=True))
        if not text:
            continue
        if node.name in {"h1", "h2", "h3"}:
            heading = text
            continue
        paragraph_number += 1
        segments.append({
            "text": text,
            "locator": {"section": heading, "paragraph": paragraph_number},
        })
    return segments


def _extract_text_file(path: Path) -> list[dict[str, Any]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return _paragraph_segments(raw, locator_prefix="lines")


def _paragraph_segments(raw: str, *, locator_prefix: str) -> list[dict[str, Any]]:
    segments = []
    line_cursor = 1
    for paragraph in re.split(r"\n\s*\n+", str(raw or "")):
        text = paragraph.strip()
        if not text:
            line_cursor += paragraph.count("\n") + 1
            continue
        line_count = text.count("\n") + 1
        segments.append({
            "text": " ".join(line.strip() for line in text.splitlines() if line.strip()),
            "locator": {
                "kind": locator_prefix,
                "line_start": line_cursor,
                "line_end": line_cursor + line_count - 1,
            },
        })
        line_cursor += line_count + 1
    return segments


def _extract_image(path: Path) -> list[dict[str, Any]]:
    text = _run_tesseract(path)
    return [
        {"text": paragraph, "locator": {"image": path.name, "paragraph": index}}
        for index, paragraph in enumerate(_split_paragraphs(text), start=1)
    ]


def _run_tesseract(path: Path) -> str:
    binary = shutil.which("tesseract")
    if not binary:
        raise ExtractionError("Tesseract is not installed")
    languages = _available_tesseract_languages(binary)
    requested = [lang for lang in ("cat", "spa", "eng", "fra") if lang in languages]
    cmd = [binary, str(path), "stdout"]
    if requested:
        cmd.extend(["-l", "+".join(requested)])
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180, check=False)
    if proc.returncode != 0:
        raise ExtractionError(proc.stderr.strip() or "Tesseract failed")
    return proc.stdout.strip()


def _available_tesseract_languages(binary: str) -> set[str]:
    try:
        proc = subprocess.run(
            [binary, "--list-langs"],
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
        return {line.strip() for line in proc.stdout.splitlines()[1:] if line.strip()}
    except Exception:
        return set()


def _extract_audio(path: Path) -> list[dict[str, Any]]:
    from backend.services.transcription import transcribe

    result = transcribe(str(path))
    return [
        {
            "text": str(item.get("text") or "").strip(),
            "locator": {
                "start": float(item.get("start") or 0),
                "end": float(item.get("end") or 0),
            },
        }
        for item in result.get("segments") or []
        if str(item.get("text") or "").strip()
    ]


def _extract_video(path: Path) -> list[dict[str, Any]]:
    segments = _extract_audio(path)
    if not shutil.which("ffmpeg") or not shutil.which("tesseract"):
        return segments
    with tempfile.TemporaryDirectory(
        prefix="gnosi-llm-wiki-frames-",
        dir=_temporary_root(),
    ) as tmp:
        pattern = str(Path(tmp) / "frame-%04d.jpg")
        cmd = [
            shutil.which("ffmpeg") or "ffmpeg",
            "-hide_banner",
            "-loglevel",
            "error",
            "-i",
            str(path),
            "-vf",
            "fps=1/60,scale=1280:-2",
            "-frames:v",
            "40",
            pattern,
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=600, check=False)
        if proc.returncode != 0:
            logger.warning("llm_wiki keyframe extraction failed: %s", proc.stderr.strip())
            return segments
        for frame_index, frame in enumerate(sorted(Path(tmp).glob("frame-*.jpg")), start=1):
            try:
                text = _run_tesseract(frame)
            except Exception:
                continue
            for paragraph in _split_paragraphs(text):
                segments.append({
                    "text": paragraph,
                    "locator": {
                        "start": float((frame_index - 1) * 60),
                        "frame": frame_index,
                        "visual": True,
                    },
                })
    return sorted(
        segments,
        key=lambda item: float((item.get("locator") or {}).get("start") or 0),
    )


def _extract_url(url: str, input_order: int) -> list[dict[str, Any]]:
    if _looks_like_streaming_page(url):
        return _extract_streaming_url(url, input_order)

    response, final_url = _download_public_url(url)
    content_type = str(response.headers.get("content-type") or "").split(";", 1)[0].lower()
    suffix = _suffix_for_response(final_url, content_type)
    if content_type in {
        "application/rss+xml",
        "application/atom+xml",
        "application/xml",
        "text/xml",
    }:
        media_url = _embedded_media_url(response.content, final_url, xml=True)
        if not media_url:
            raise ExtractionError("The feed did not contain a public media enclosure")
        origins = _extract_url(media_url, input_order)
        for origin in origins:
            origin.setdefault("aliases", []).append({
                "kind": "feed",
                "label": final_url,
                "source_url": final_url,
                "input_order": input_order,
            })
        return origins
    if content_type in {"text/html", "application/xhtml+xml"} or suffix in HTML_EXTENSIONS:
        raw_html = response.content.decode(response.encoding or "utf-8", errors="replace")
        segments = _extract_html(raw_html)
        if not segments:
            try:
                return _extract_streaming_url(final_url, input_order)
            except Exception as exc:
                raise ExtractionError(
                    "The web page did not contain extractable article text or public media"
                ) from exc
        origins = [_finalize_origin({
            "kind": "url",
            "label": final_url,
            "source_url": final_url,
            "input_order": input_order,
            "segments": segments,
        })]
        media_url = _embedded_media_url(response.content, final_url)
        if media_url and media_url != final_url:
            try:
                origins.extend(_extract_url(media_url, input_order))
            except Exception as exc:  # noqa: BLE001
                logger.warning("llm_wiki embedded public media extraction failed: %s", exc)
        return origins

    with tempfile.NamedTemporaryFile(
        suffix=suffix or ".bin",
        delete=False,
        dir=_temporary_root(),
    ) as tmp:
        tmp.write(response.content)
        tmp_path = Path(tmp.name)
    try:
        origins = _extract_path(tmp_path, input_order, label=Path(urlparse(final_url).path).name or final_url)
        for origin in origins:
            origin["source_url"] = final_url
            origin["origin_id"] = _origin_id(origin)
        return origins
    finally:
        tmp_path.unlink(missing_ok=True)


def _download_public_url(url: str) -> tuple[requests.Response, str]:
    from backend.agent.web_context import is_public_http_url

    current = str(url).strip()
    for _ in range(6):
        ok, reason = is_public_http_url(current)
        if not ok:
            raise ExtractionError(f"Unsafe URL blocked: {reason}")
        response = requests.get(
            current,
            timeout=_HTTP_TIMEOUT,
            headers={"User-Agent": _USER_AGENT, "Accept-Language": "ca,es,en,fr"},
            allow_redirects=False,
            stream=True,
        )
        if response.is_redirect and response.headers.get("location"):
            current = urljoin(current, response.headers["location"])
            continue
        response.raise_for_status()
        declared = int(response.headers.get("content-length") or 0)
        if declared and declared > _MAX_DOWNLOAD_BYTES:
            raise ExtractionError("Remote source exceeds the configured download limit")
        chunks = []
        total = 0
        for chunk in response.iter_content(chunk_size=1024 * 1024):
            if not chunk:
                continue
            total += len(chunk)
            if total > _MAX_DOWNLOAD_BYTES:
                raise ExtractionError("Remote source exceeds the configured download limit")
            chunks.append(chunk)
        response._content = b"".join(chunks)  # requests exposes content through this cache.
        return response, current
    raise ExtractionError("Too many URL redirects")


def _looks_like_streaming_page(url: str) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return any(
        marker in host
        for marker in ("youtube.com", "youtu.be", "vimeo.com", "soundcloud.com")
    )


def _extract_streaming_url(url: str, input_order: int) -> list[dict[str, Any]]:
    try:
        import yt_dlp
    except Exception as exc:
        raise ExtractionError("yt-dlp is not installed") from exc
    with tempfile.TemporaryDirectory(
        prefix="gnosi-llm-wiki-stream-",
        dir=_temporary_root(),
    ) as tmp:
        template = str(Path(tmp) / "source.%(ext)s")
        options = {
            "format": "bestaudio/best",
            "outtmpl": template,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
            "max_filesize": _MAX_DOWNLOAD_BYTES,
            "restrictfilenames": True,
        }
        with yt_dlp.YoutubeDL(options) as ydl:
            info = ydl.extract_info(url, download=True)
            path = Path(ydl.prepare_filename(info))
        if not path.exists():
            candidates = list(Path(tmp).glob("source.*"))
            path = candidates[0] if candidates else path
        if not path.exists():
            raise ExtractionError("The streaming source could not be downloaded")
        segments = _extract_audio(path)
        if not segments:
            raise ExtractionError("The streaming source did not produce a transcript")
        return [_finalize_origin({
            "kind": "stream",
            "label": str(info.get("title") or url),
            "source_url": str(info.get("webpage_url") or url),
            "input_order": input_order,
            "segments": segments,
        })]


def _embedded_media_url(content: bytes, base_url: str, *, xml: bool = False) -> str:
    """Return the first public audio/video enclosure referenced by a page."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "xml" if xml else "html.parser")
    candidates = []
    if xml:
        for node in soup.find_all(["enclosure", "content"]):
            media_type = str(node.get("type") or "").lower()
            if media_type.startswith(("audio/", "video/")) or node.name == "enclosure":
                candidates.append(node.get("url"))
    else:
        for node in soup.find_all(["audio", "video", "source"]):
            candidates.append(node.get("src"))
        for node in soup.find_all("meta"):
            prop = str(node.get("property") or node.get("name") or "").lower()
            if prop in {"og:audio", "og:audio:url", "og:video", "og:video:url"}:
                candidates.append(node.get("content"))
    for candidate in candidates:
        value = str(candidate or "").strip()
        if value:
            return urljoin(base_url, value)
    return ""


def _temporary_root() -> Path:
    """Keep extraction temporaries in local Gnosi data, never in the vault."""
    from backend.api.vault_routes import get_p

    root = get_p("LOCAL_DATA") / "llm_wiki" / "tmp"
    root.mkdir(parents=True, exist_ok=True)
    return root


def _suffix_for_response(url: str, content_type: str) -> str:
    suffix = Path(urlparse(url).path).suffix.lower()
    if suffix:
        return suffix
    guessed = mimetypes.guess_extension(content_type or "") or ""
    return ".jpg" if guessed == ".jpe" else guessed


def _split_paragraphs(text: str) -> list[str]:
    raw = str(text or "").strip()
    if not raw:
        return []
    paragraphs = [
        " ".join(piece.split())
        for piece in re.split(r"\n\s*\n+|(?<=\.)\s*\n+", raw)
        if " ".join(piece.split())
    ]
    return paragraphs


def _finalize_origin(origin: dict[str, Any]) -> dict[str, Any]:
    segments = []
    for order, segment in enumerate(origin.get("segments") or [], start=1):
        text = " ".join(str(segment.get("text") or "").split()).strip()
        if not text:
            continue
        segments.append({
            "id": "",
            "order": order,
            "text": text,
            "locator": segment.get("locator") or {},
        })
    content_hash = hashlib.sha256(
        "\n".join(segment["text"] for segment in segments).encode("utf-8")
    ).hexdigest()
    origin = {
        **origin,
        "content_hash": content_hash,
        "segments": segments,
        "aliases": origin.get("aliases") or [],
    }
    origin["origin_id"] = _origin_id(origin)
    for segment in origin["segments"]:
        short = hashlib.sha256(segment["text"].encode("utf-8")).hexdigest()[:8]
        segment["id"] = f"{origin['origin_id']}-s{segment['order']}-{short}"
    return origin


def _origin_id(origin: dict[str, Any]) -> str:
    value = json.dumps(
        {
            "kind": origin.get("kind"),
            "label": origin.get("label"),
            "url": origin.get("source_url"),
            "hash": origin.get("content_hash"),
        },
        sort_keys=True,
        ensure_ascii=False,
    )
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:16]


def _deduplicate_origins(origins: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    unique: list[dict[str, Any]] = []
    by_hash: dict[str, dict[str, Any]] = {}
    for origin in origins:
        content_hash = str(origin.get("content_hash") or "")
        existing = by_hash.get(content_hash)
        if existing is not None:
            existing.setdefault("aliases", []).append({
                "kind": origin.get("kind"),
                "label": origin.get("label"),
                "source_url": origin.get("source_url"),
                "input_order": origin.get("input_order"),
            })
            continue
        by_hash[content_hash] = origin
        unique.append(origin)
    return unique
